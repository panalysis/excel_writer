

import openpyxl
from docx import Document
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE

doc_items = {
    'user_stories': [],
    'reporting_requirements': []
}

class UserStory:

    def __init__(self, id=0, story=None, reqs=None, priority=None):
        self.id =id
        self.description = story
        self.reqs = reqs
        self.priority = priority

    def __str__(self):

        mystr = '''User Story %s
        %s
        Requirements: %s
        Priority: %s
        '''
        return mystr % (str(self.id),self.description,self.reqs,self.priority)

class ReportingRequirement:

    def __init__(self):
        self.id = ""
        self.user_story_id = ""
        self.title = ""
        self.description = ""
        self.purpose = ""
        self.example_url = ""
        self.metrics = ""
        self.dimensions = ""
        self.aggregation = ""
        self.segment = ""
        self.charting = ""
        self.report_users = ""
        self.report_frequency = ""
        self.data_source = ""
        self.standard_ga = ""
        self.custom_report = ""
        self.priority = ""


    def __str__(self):
        return str(self.title)

def load_user_stories(d):
    s1 = d.get_sheet_by_name('User Stories')
    in_data = False
    headings = {}
    num_headings = 4

    for r in s1.rows:

        if in_data:
            my_story = UserStory()
            my_story.id = r[headings['ID']].value
            my_story.description = r[headings['User Story']].value
            my_story.reqs = r[headings['Reporting Requirements']].value
            my_story.priority = r[headings['Priority']].value

            doc_items['user_stories'].append(my_story)

        if in_data is False:
            in_data, headings = check_header(num_headings, r)

def load_reporting_requirements(d):
    s1 = d.get_sheet_by_name('Reporting Requirements')
    in_data = False
    headings = {}
    num_headings = 20

    for r in s1.rows:

        if in_data:
            my_req = ReportingRequirement()
            my_req.id = r[headings['ID']].value
            my_req.user_story_id = r[headings['User Story']].value
            my_req.title = r[headings['Title']].value
            my_req.description = r[headings['Description']].value
            my_req.purpose = r[headings['Purpose of Report']].value
            my_req.example_url = r[headings['URL']].value
            my_req.metrics = r[headings['Metrics']].value
            my_req.dimensions = r[headings['Dimensions']].value
            my_req.aggregation = r[headings['Aggregation']].value
            my_req.segment = r[headings['Segment']].value
            my_req.charting = r[headings['Charting']].value
            my_req.report_users = r[headings['Report Users']].value
            my_req.report_frequency = r[headings['Report Frequency']].value
            my_req.standard_ga = r[headings['Standard GA Report Functionality']].value
            my_req.custom_report = r[headings['Requires Custom Report in Google Analytics']].value
            my_req.priority = r[headings['Priority']].value

            doc_items['reporting_requirements'].append(my_req)

        if in_data is False:
            in_data, headings = check_header(num_headings, r)


def check_header(num_headings, r):
    np = 0
    headings = {}
    in_data = False
    for i in range(num_headings):
        if r[i].value is not None:
            np += 1
    if np >= num_headings:
        in_data = True
        i = 0
        for a in r:
            headings[a.value] = i
            i += 1
    return in_data, headings


def write_report(file_name):

    document = Document()
    document.add_heading('User Stories', level=1)

    styles = document.styles
    us_style = styles.add_style("Panalysis User Stories", WD_STYLE_TYPE.TABLE)
    us_style.base_style = styles["Medium Grid 1 Accent 1"]

    rr_style = styles.add_style("Panalysis Reporting Requirement", WD_STYLE_TYPE.TABLE)
    rr_style.base_style = styles["Medium List 2 Accent 1"]

    table1 = document.add_table(rows=1, cols=4, style=us_style)

    table1.columns[0].width = Cm(3)
    table1.columns[1].width = Cm(12)
    table1.columns[2].width = Cm(5)
    table1.columns[3].width = Cm(4)

    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Number'
    hdr_cells[1].text = 'Details'
    hdr_cells[2].text = 'Priority'
    hdr_cells[3].text = 'Reporting Requirements'

    for i in doc_items['user_stories']:

        row_cells = table1.add_row().cells
        row_cells[0].text = str(i.id)
        row_cells[1].text = str(i.description)
        row_cells[2].text = str(i.priority)
        row_cells[3].text = str(i.reqs)


    document.add_page_break()

    document.add_heading('Reporting Requirements', level=1)

    for i in doc_items['reporting_requirements']:

        document.add_heading(str(i.id) + ": " +  i.title, level=2)
        table2 = document.add_table(rows=1, cols=2, style=rr_style)

        table2.columns[0].width = Cm(2)
        table2.columns[1].width = Cm(15)

        requirements_table = ["Description","Purpose","User Stories","Priority","Data Source","Dimensions","Metrics","Segments","Aggregation","Requires Custom Report"]
        for x in requirements_table:
            row_cells = table2.add_row().cells
            row_cells[0].text = str(x) + ":"

            if x == "Description":
                row_cells[1].text = i.description
            elif x == "Purpose":
                row_cells[1].text = i.purpose
            elif x == "User Stories":
                row_cells[1].text = str(i.user_story_id)
            elif x == "Priority":
                row_cells[1].text = i.priority
            elif x == "Data Source":
                row_cells[1].text = i.data_source
            elif x == "Dimensions":
                row_cells[1].text = i.dimensions
            elif x == "Metrics":
                row_cells[1].text = i.metrics
            elif x == "Segments":
                row_cells[1].text = i.segment
            elif x == "Aggregation":
                row_cells[1].text = i.aggregation
            elif x == "Requires Custom Report":
                row_cells[1].text = i.custom_report


    document.save(file_name)


def main():
    d = openpyxl.load_workbook("test.xlsx")
    load_user_stories(d)
    load_reporting_requirements(d)
    write_report("test.docx")

    for i in doc_items['user_stories']:
        print(i)

    for i in doc_items['reporting_requirements']:
        print(i)

if __name__ == '__main__':
    main()